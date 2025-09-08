#!/usr/bin/env python3
"""
Script para crear documentación completa de la aplicación Inmuebles
Genera un documento Word con estructura, screenshots y diagramas
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os
import json

def create_web_structure_documentation():
    """Crea el documento Word con la documentación completa"""
    
    # Crear nuevo documento
    doc = Document()
    
    # Configurar estilos
    setup_styles(doc)
    
    # 1. PORTADA
    create_title_page(doc)
    
    # 2. ÍNDICE (placeholder)
    create_table_of_contents(doc)
    
    # 3. INTRODUCCIÓN Y OVERVIEW
    create_introduction(doc)
    
    # 4. ARQUITECTURA DEL SISTEMA
    create_architecture_section(doc)
    
    # 5. ESTRUCTURA DE NAVEGACIÓN
    create_navigation_structure(doc)
    
    # 6. MÓDULOS PRINCIPALES
    create_main_modules_section(doc)
    
    # 7. PÁGINAS Y FUNCIONALIDADES
    create_pages_documentation(doc)
    
    # 8. DIAGRAMAS DE FLUJO
    create_flow_diagrams(doc)
    
    # 9. CONFIGURACIÓN Y DEPLOYMENT
    create_config_section(doc)
    
    # 10. CONCLUSIONES
    create_conclusions(doc)
    
    # Guardar documento
    doc_path = r"C:\Users\davsa\inmuebles\backend\Documentacion_Aplicacion_Inmuebles.docx"
    doc.save(doc_path)
    print(f"Documento creado exitosamente en: {doc_path}")
    
    return doc_path

def setup_styles(doc):
    """Configurar estilos personalizados"""
    
    # Estilo para títulos principales
    styles = doc.styles
    
    # Estilo título 1
    if 'Titulo Principal' not in styles:
        title_style = styles.add_style('Titulo Principal', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

def create_title_page(doc):
    """Crear página de título"""
    
    title = doc.add_heading('DOCUMENTACIÓN TÉCNICA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('APLICACIÓN DE GESTIÓN DE INMUEBLES', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n')
    
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run('Sistema Completo de Financial Agent\n').font.size = Pt(14)
    info_para.add_run('Gestión de Propiedades, Movimientos Financieros y Analytics\n').font.size = Pt(12)
    
    doc.add_paragraph('\n\n')
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run('Septiembre 2025').font.size = Pt(12)
    
    doc.add_page_break()

def create_table_of_contents(doc):
    """Crear tabla de contenidos"""
    
    doc.add_heading('ÍNDICE', level=1)
    
    toc_items = [
        "1. Introducción y Overview",
        "2. Arquitectura del Sistema",
        "3. Estructura de Navegación",
        "4. Módulos Principales",
        "5. Páginas y Funcionalidades",
        "6. Diagramas de Flujo",
        "7. Configuración y Deployment",
        "8. Conclusiones"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()

def create_introduction(doc):
    """Crear sección de introducción"""
    
    doc.add_heading('1. INTRODUCCIÓN Y OVERVIEW', level=1)
    
    doc.add_paragraph(
        'La Aplicación de Gestión de Inmuebles es una solución completa desarrollada con '
        'tecnologías modernas para la administración integral de propiedades inmobiliarias. '
        'El sistema permite gestionar múltiples propiedades, sus movimientos financieros, '
        'contratos de alquiler, hipotecas y generar análisis detallados de rentabilidad.'
    )
    
    doc.add_heading('1.1 Características Principales', level=2)
    
    features = [
        '🏠 Gestión completa de propiedades inmobiliarias',
        '💰 Tracking de movimientos financieros automático',
        '📊 Analytics y reportes de rentabilidad (ROI, Cash Flow)',
        '🏦 Integración con entidades bancarias (Bankinter)',
        '📋 Gestión de contratos de alquiler',
        '🏦 Calculadora y gestión de hipotecas',
        '📈 Seguimiento de Euribor para hipotecas variables',
        '🤖 Clasificación inteligente de movimientos',
        '📱 Interfaz responsive y moderna',
        '🔐 Sistema de autenticación y autorización'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('1.2 Stack Tecnológico', level=2)
    
    doc.add_paragraph('Frontend:')
    frontend_tech = [
        'Next.js 15.4.6 con App Router',
        'React 18 con TypeScript',
        'Tailwind CSS para estilos',
        'Turbopack para desarrollo rápido'
    ]
    
    for tech in frontend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_paragraph('Backend:')
    backend_tech = [
        'FastAPI con Python',
        'SQLModel/SQLAlchemy para ORM',
        'PostgreSQL/SQLite para base de datos',
        'JWT para autenticación'
    ]
    
    for tech in backend_tech:
        doc.add_paragraph(tech, style='List Bullet')

def create_architecture_section(doc):
    """Crear sección de arquitectura"""
    
    doc.add_heading('2. ARQUITECTURA DEL SISTEMA', level=1)
    
    doc.add_paragraph(
        'El sistema utiliza una arquitectura de microservicios con separación clara entre '
        'frontend y backend, desplegados en plataformas cloud diferentes para optimización de costos y rendimiento.'
    )
    
    doc.add_heading('2.1 Diagrama de Arquitectura', level=2)
    
    # Placeholder para diagrama
    doc.add_paragraph('📊 [DIAGRAMA DE ARQUITECTURA - A INSERTAR]')
    doc.add_paragraph('')
    
    doc.add_paragraph('Componentes principales:')
    components = [
        'Frontend (Next.js) - Desplegado en Vercel',
        'Backend API (FastAPI) - Desplegado en Render/Railway',
        'Base de Datos - PostgreSQL en la nube',
        'Integración Bancaria - APIs de Bankinter PSD2',
        'Sistema de Archivos - Almacenamiento de documentos'
    ]
    
    for comp in components:
        doc.add_paragraph(comp, style='List Bullet')
    
    doc.add_heading('2.2 Flujo de Datos', level=2)
    
    doc.add_paragraph(
        'El flujo de datos sigue un patrón RESTful donde el frontend consume APIs del backend '
        'que a su vez interactúa con la base de datos y servicios externos como la integración bancaria.'
    )

def create_navigation_structure(doc):
    """Crear estructura de navegación"""
    
    doc.add_heading('3. ESTRUCTURA DE NAVEGACIÓN', level=1)
    
    doc.add_paragraph(
        'La aplicación está organizada en módulos principales accesibles desde el menú de navegación. '
        'La estructura de rutas utiliza el App Router de Next.js 13+.'
    )
    
    # Estructura de rutas
    route_structure = {
        "/": "Página principal/Dashboard",
        "/login": "Página de autenticación",
        "/dashboard": "Dashboard principal",
        "/dashboard/properties": "Lista de propiedades",
        "/financial-agent": "Hub del agente financiero",
        "/financial-agent/movements": "Gestión de movimientos financieros",
        "/financial-agent/analytics": "Dashboard de analytics y métricas",
        "/financial-agent/contracts": "Gestión de contratos de alquiler",
        "/financial-agent/rules": "Reglas de clasificación automática",
        "/financial-agent/integrations": "Integraciones bancarias",
        "/financial-agent/integrations/bankinter": "Configuración Bankinter",
        "/financial-agent/mortgage-calculator": "Calculadora de hipotecas",
        "/financial-agent/euribor": "Gestión de tipos Euribor",
        "/financial-agent/documents": "Gestor de documentos",
        "/financial-agent/notifications": "Centro de notificaciones",
        "/financial-agent/tax-assistant": "Asistente fiscal",
        "/financial-agent/smart-classifier": "Clasificador inteligente",
        "/financial-agent/property/[id]": "Vista detallada de propiedad",
        "/financial-agent/property/[id]/reports": "Informes financieros de propiedad",
        "/financial-agent/property/[id]/mortgage": "Gestión hipoteca de propiedad",
        "/financial-agent/property/[id]/rules": "Reglas específicas de propiedad",
        "/financial-agent/property/[id]/contracts/new": "Crear nuevo contrato"
    }
    
    doc.add_heading('3.1 Mapa de Rutas', level=2)
    
    for route, description in route_structure.items():
        p = doc.add_paragraph()
        p.add_run(f"{route}").font.bold = True
        p.add_run(f" - {description}")

def create_main_modules_section(doc):
    """Crear sección de módulos principales"""
    
    doc.add_heading('4. MÓDULOS PRINCIPALES', level=1)
    
    modules = [
        {
            'name': '🏠 Dashboard de Propiedades',
            'description': 'Gestión central de todas las propiedades inmobiliarias',
            'features': [
                'Lista completa de propiedades',
                'Información básica (dirección, tipo, m², habitaciones)',
                'Estados financieros resumidos',
                'Navegación rápida a detalles'
            ]
        },
        {
            'name': '💰 Financial Agent',
            'description': 'Núcleo del sistema de gestión financiera',
            'features': [
                'Dashboard central con métricas principales',
                'Resumen de todas las propiedades',
                'Acceso rápido a funcionalidades',
                'Indicadores de rendimiento'
            ]
        },
        {
            'name': '📊 Movimientos Financieros',
            'description': 'Gestión completa de ingresos y gastos',
            'features': [
                'Lista filtrable de movimientos',
                'Clasificación automática e inteligente',
                'Asignación a propiedades',
                'Importación masiva desde bancos',
                'Edición y eliminación de movimientos'
            ]
        },
        {
            'name': '📈 Analytics y Reportes',
            'description': 'Sistema avanzado de análisis financiero',
            'features': [
                'Dashboard con KPIs principales',
                'Gráficos de rentabilidad por propiedad',
                'Análisis de cash flow temporal',
                'Métricas de ROI y Cap Rate',
                'Comparativas entre propiedades'
            ]
        },
        {
            'name': '📋 Gestión de Contratos',
            'description': 'Administración de contratos de alquiler',
            'features': [
                'Registro de contratos por propiedad',
                'Seguimiento de inquilinos',
                'Gestión de fianzas y depósitos',
                'Renovaciones y finalizaciones'
            ]
        },
        {
            'name': '🏦 Sistema Hipotecario',
            'description': 'Gestión integral de hipotecas',
            'features': [
                'Registro de datos hipotecarios',
                'Calculadora de cuotas',
                'Seguimiento de revisiones Euribor',
                'Cronograma de pagos',
                'Análisis de impacto en rentabilidad'
            ]
        },
        {
            'name': '🔗 Integraciones Bancarias',
            'description': 'Conexión con entidades financieras',
            'features': [
                'Integración PSD2 con Bankinter',
                'Descarga automática de movimientos',
                'Sincronización programada',
                'Mapeo automático de transacciones'
            ]
        }
    ]
    
    for module in modules:
        doc.add_heading(f'4.1 {module["name"]}', level=2)
        doc.add_paragraph(module['description'])
        
        doc.add_paragraph('Características principales:')
        for feature in module['features']:
            doc.add_paragraph(feature, style='List Bullet')
        
        doc.add_paragraph('📷 [SCREENSHOT PLACEHOLDER]')
        doc.add_paragraph('')

def create_pages_documentation(doc):
    """Crear documentación detallada de páginas"""
    
    doc.add_heading('5. PÁGINAS Y FUNCIONALIDADES DETALLADAS', level=1)
    
    pages_info = [
        {
            'title': '🔐 Página de Login',
            'route': '/login',
            'description': 'Sistema de autenticación con JWT',
            'features': [
                'Formulario de login con validación',
                'Manejo de errores de autenticación',
                'Redirección automática post-login',
                'Diseño responsive'
            ]
        },
        {
            'title': '🏠 Dashboard Principal',
            'route': '/financial-agent',
            'description': 'Hub central de la aplicación',
            'features': [
                'Resumen de propiedades',
                'Métricas financieras generales por año',
                'Navegación a módulos principales',
                'Filtros por año'
            ]
        },
        {
            'title': '💰 Gestión de Movimientos',
            'route': '/financial-agent/movements',
            'description': 'CRUD completo de movimientos financieros',
            'features': [
                'Lista paginada con filtros avanzados',
                'Clasificación automática e inteligente',
                'Asignación/reasignación a propiedades',
                'Importación desde CSV/Excel',
                'Sistema de fallback para eliminación',
                'Vista de movimientos sin asignar'
            ]
        },
        {
            'title': '📊 Analytics Dashboard',
            'route': '/financial-agent/analytics',
            'description': 'Centro de análisis financiero',
            'features': [
                'KPIs principales del portfolio',
                'Gráficos de rentabilidad',
                'Análisis temporal de cash flow',
                'Comparativas entre propiedades',
                'Métricas de ROI consolidadas'
            ]
        },
        {
            'title': '🏠 Vista de Propiedad',
            'route': '/financial-agent/property/[id]',
            'description': 'Dashboard específico de cada propiedad',
            'features': [
                'Información básica de la propiedad',
                'Resumen financiero anual',
                'Navegación a subsecciones',
                'Edición de datos básicos'
            ]
        },
        {
            'title': '📈 Informes de Propiedad',
            'route': '/financial-agent/property/[id]/reports',
            'description': 'Reportes financieros detallados por propiedad',
            'features': [
                'Cash flow mensual visualizado',
                'Desglose por categorías',
                'Métricas clave (ROI, promedios)',
                'Selector de año',
                'Función de impresión/exportación',
                'Cálculo inteligente de promedios mensuales'
            ]
        },
        {
            'title': '🏦 Gestión Hipotecaria',
            'route': '/financial-agent/property/[id]/mortgage',
            'description': 'Administración completa de hipotecas',
            'features': [
                'Registro de datos hipotecarios',
                'Gestión de revisiones Euribor',
                'Cronograma de amortización',
                'Calculadora integrada',
                'Análisis ROI (deshabilitado temporalmente)'
            ]
        },
        {
            'title': '📋 Contratos de Alquiler',
            'route': '/financial-agent/contracts',
            'description': 'Gestión de contratos y inquilinos',
            'features': [
                'Lista de contratos activos/vencidos',
                'Información de inquilinos',
                'Seguimiento de fianzas',
                'Gestión de renovaciones'
            ]
        },
        {
            'title': '🤖 Reglas de Clasificación',
            'route': '/financial-agent/rules',
            'description': 'Sistema de automatización para movimientos',
            'features': [
                'Creación de reglas personalizadas',
                'Patrones de texto para clasificación',
                'Asignación automática de categorías',
                'Gestión de excepciones'
            ]
        },
        {
            'title': '🔗 Integraciones',
            'route': '/financial-agent/integrations',
            'description': 'Centro de integraciones externas',
            'features': [
                'Lista de integraciones disponibles',
                'Estado de conexiones',
                'Configuración de credenciales',
                'Logs de sincronización'
            ]
        },
        {
            'title': '🏦 Integración Bankinter',
            'route': '/financial-agent/integrations/bankinter',
            'description': 'Configuración específica para Bankinter PSD2',
            'features': [
                'Configuración de credenciales',
                'Test de conexión',
                'Sincronización manual/automática',
                'Mapeo de cuentas'
            ]
        },
        {
            'title': '🧮 Calculadora Hipotecaria',
            'route': '/financial-agent/mortgage-calculator',
            'description': 'Herramienta de cálculo hipotecario',
            'features': [
                'Cálculo de cuotas variables/fijas',
                'Simulación con diferentes Euribor',
                'Cronograma de amortización',
                'Comparativa de opciones'
            ]
        },
        {
            'title': '📈 Gestión Euribor',
            'route': '/financial-agent/euribor',
            'description': 'Seguimiento de tipos de interés',
            'features': [
                'Histórico de tipos Euribor',
                'Actualización manual/automática',
                'Impacto en hipotecas existentes',
                'Proyecciones futuras'
            ]
        }
    ]
    
    for page in pages_info:
        doc.add_heading(f'5.{pages_info.index(page)+1} {page["title"]}', level=2)
        
        p = doc.add_paragraph()
        p.add_run('Ruta: ').font.bold = True
        p.add_run(page['route'])
        
        doc.add_paragraph(page['description'])
        
        doc.add_paragraph('Funcionalidades:')
        for feature in page['features']:
            doc.add_paragraph(feature, style='List Bullet')
        
        doc.add_paragraph('📷 [SCREENSHOT PLACEHOLDER - ' + page['title'] + ']')
        doc.add_paragraph('')

def create_flow_diagrams(doc):
    """Crear sección de diagramas de flujo"""
    
    doc.add_heading('6. DIAGRAMAS DE FLUJO PRINCIPALES', level=1)
    
    doc.add_heading('6.1 Flujo de Autenticación', level=2)
    doc.add_paragraph('📊 [DIAGRAMA - Proceso de login y autorización]')
    
    auth_flow = [
        'Usuario accede a aplicación',
        'Redirección a /login si no está autenticado',
        'Validación de credenciales en backend',
        'Generación de token JWT',
        'Almacenamiento en localStorage',
        'Redirección a dashboard'
    ]
    
    for i, step in enumerate(auth_flow):
        doc.add_paragraph(f'{i+1}. {step}')
    
    doc.add_paragraph('')
    
    doc.add_heading('6.2 Flujo de Gestión de Movimientos', level=2)
    doc.add_paragraph('📊 [DIAGRAMA - Proceso de importación y clasificación]')
    
    movement_flow = [
        'Importación desde banco o manual',
        'Análisis con clasificador inteligente',
        'Aplicación de reglas de clasificación',
        'Asignación a propiedad (si aplica)',
        'Categorización (Renta, Hipoteca, etc.)',
        'Almacenamiento en base de datos',
        'Actualización de métricas en tiempo real'
    ]
    
    for i, step in enumerate(movement_flow):
        doc.add_paragraph(f'{i+1}. {step}')
    
    doc.add_paragraph('')
    
    doc.add_heading('6.3 Flujo de Análisis Financiero', level=2)
    doc.add_paragraph('📊 [DIAGRAMA - Proceso de cálculo de métricas]')
    
    analytics_flow = [
        'Recopilación de movimientos por período',
        'Cálculo de ingresos y gastos totales',
        'Determinación de cash flow neto',
        'Cálculo de ROI y Cap Rate',
        'Análisis de tendencias temporales',
        'Generación de gráficos y reportes',
        'Comparativas entre propiedades'
    ]
    
    for i, step in enumerate(analytics_flow):
        doc.add_paragraph(f'{i+1}. {step}')

def create_config_section(doc):
    """Crear sección de configuración"""
    
    doc.add_heading('7. CONFIGURACIÓN Y DEPLOYMENT', level=1)
    
    doc.add_heading('7.1 Variables de Entorno', level=2)
    
    env_vars = [
        'NEXT_PUBLIC_API_URL - URL del backend API',
        'DATABASE_URL - Conexión a base de datos',
        'JWT_SECRET - Clave para tokens JWT',
        'BANKINTER_CLIENT_ID - Credenciales PSD2',
        'BANKINTER_CLIENT_SECRET - Credenciales PSD2'
    ]
    
    for var in env_vars:
        doc.add_paragraph(var, style='List Bullet')
    
    doc.add_heading('7.2 Proceso de Deployment', level=2)
    
    doc.add_paragraph('Frontend (Vercel):')
    frontend_deploy = [
        'Build automático desde Git',
        'Optimización de assets',
        'CDN global',
        'SSL automático'
    ]
    
    for step in frontend_deploy:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_paragraph('Backend (Render/Railway):')
    backend_deploy = [
        'Deployment automático',
        'Base de datos PostgreSQL',
        'Scaling automático',
        'Health checks'
    ]
    
    for step in backend_deploy:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_heading('7.3 Scripts de Deployment', level=2)
    
    scripts = [
        'deploy-inmuebles-vercel.bat - Deploy frontend',
        'deploy-backend-fixed.bat - Deploy backend',
        'check-deployment.bat - Verificar deployments'
    ]
    
    for script in scripts:
        doc.add_paragraph(script, style='List Bullet')

def create_conclusions(doc):
    """Crear sección de conclusiones"""
    
    doc.add_heading('8. CONCLUSIONES Y PRÓXIMOS PASOS', level=1)
    
    doc.add_paragraph(
        'La aplicación de gestión de inmuebles representa una solución completa y moderna '
        'para la administración de propiedades inmobiliarias. El sistema integra múltiples '
        'funcionalidades en una interfaz coherente y fácil de usar.'
    )
    
    doc.add_heading('8.1 Fortalezas del Sistema', level=2)
    
    strengths = [
        '✅ Arquitectura moderna y escalable',
        '✅ Integración bancaria automatizada',
        '✅ Sistema de clasificación inteligente',
        '✅ Analytics avanzados de rentabilidad',
        '✅ Interfaz responsive y intuitiva',
        '✅ Deployment automatizado',
        '✅ Gestión completa del ciclo de vida financiero'
    ]
    
    for strength in strengths:
        doc.add_paragraph(strength)
    
    doc.add_heading('8.2 Áreas de Mejora Identificadas', level=2)
    
    improvements = [
        '🔄 Implementación de notificaciones push',
        '📱 Desarrollo de app móvil nativa',
        '🤖 Expansión del sistema de IA',
        '📊 Más integraciones bancarias',
        '🔐 Autenticación de dos factores',
        '☁️ Backup automático en cloud'
    ]
    
    for improvement in improvements:
        doc.add_paragraph(improvement)
    
    doc.add_paragraph('\n---\n')
    doc.add_paragraph('Documento generado automáticamente - Septiembre 2025')
    
    return doc

if __name__ == "__main__":
    try:
        print("Creando documentacion de la aplicacion...")
        doc_path = create_web_structure_documentation()
        print("Documentacion creada exitosamente!")
        print(f"Ubicacion: {doc_path}")
    except Exception as e:
        print(f"Error al crear documentacion: {str(e)}")