#!/usr/bin/env python3
"""
Script simplificado para generar screenshots usando playwright
Como alternativa al script con Selenium
"""

import os
import asyncio

def create_screenshot_placeholders():
    """Crear placeholders de screenshots y actualizar documentación"""
    
    # Crear directorio de screenshots
    screenshot_dir = r"C:\Users\davsa\inmuebles\backend\screenshots"
    os.makedirs(screenshot_dir, exist_ok=True)
    
    # Lista de páginas a documentar
    pages_info = [
        {
            'filename': '01_login.png',
            'title': 'Página de Login',
            'description': 'Sistema de autenticación con formulario de usuario y contraseña'
        },
        {
            'filename': '02_dashboard_principal.png', 
            'title': 'Dashboard Principal - Financial Agent',
            'description': 'Hub central con resumen de propiedades y métricas por año'
        },
        {
            'filename': '03_gestion_movimientos.png',
            'title': 'Gestión de Movimientos Financieros',
            'description': 'CRUD completo con filtros, clasificación y asignación a propiedades'
        },
        {
            'filename': '04_analytics_dashboard.png',
            'title': 'Dashboard de Analytics',
            'description': 'KPIs, gráficos de rentabilidad y análisis comparativo'
        },
        {
            'filename': '05_gestion_contratos.png',
            'title': 'Gestión de Contratos',
            'description': 'Administración de contratos de alquiler e inquilinos'
        },
        {
            'filename': '06_reglas_clasificacion.png',
            'title': 'Reglas de Clasificación',
            'description': 'Sistema de automatización para categorización de movimientos'
        },
        {
            'filename': '07_integraciones.png',
            'title': 'Centro de Integraciones',
            'description': 'Gestión de conexiones con servicios externos'
        },
        {
            'filename': '08_integracion_bankinter.png',
            'title': 'Integración Bankinter',
            'description': 'Configuración PSD2 y sincronización automática'
        },
        {
            'filename': '09_calculadora_hipoteca.png',
            'title': 'Calculadora Hipotecaria',
            'description': 'Herramienta de simulación y cálculo de cuotas'
        },
        {
            'filename': '10_gestion_euribor.png',
            'title': 'Gestión de Euribor',
            'description': 'Seguimiento histórico de tipos de interés'
        },
        {
            'filename': '11_gestor_documentos.png',
            'title': 'Gestor de Documentos',
            'description': 'Sistema de archivos organizados por propiedad'
        },
        {
            'filename': '12_notificaciones.png',
            'title': 'Centro de Notificaciones',
            'description': 'Alertas, recordatorios y mensajes del sistema'
        },
        {
            'filename': '13_asistente_fiscal.png',
            'title': 'Asistente Fiscal',
            'description': 'Herramientas para declaraciones y optimización fiscal'
        },
        {
            'filename': '14_clasificador_inteligente.png',
            'title': 'Clasificador Inteligente',
            'description': 'Sistema de IA para clasificación automática'
        },
        {
            'filename': '15_propiedad_vista_general.png',
            'title': 'Vista General de Propiedad',
            'description': 'Dashboard específico con métricas de la propiedad'
        },
        {
            'filename': '16_informes_propiedad.png',
            'title': 'Informes Financieros de Propiedad',
            'description': 'Reportes detallados con cash flow y análisis mensual'
        },
        {
            'filename': '17_hipoteca_propiedad.png',
            'title': 'Gestión Hipotecaria de Propiedad',
            'description': 'Administración de hipoteca con cronograma y revisiones'
        },
        {
            'filename': '18_reglas_propiedad.png',
            'title': 'Reglas Específicas de Propiedad',
            'description': 'Configuración de reglas particulares para la propiedad'
        }
    ]
    
    # Crear archivo de información sobre screenshots
    info_content = """# INFORMACIÓN SOBRE SCREENSHOTS DE LA APLICACIÓN

Este directorio debería contener screenshots automáticos de todas las páginas principales.

## Páginas documentadas:

"""
    
    for page in pages_info:
        info_content += f"### {page['title']}\n"
        info_content += f"**Archivo:** {page['filename']}\n"
        info_content += f"**Descripción:** {page['description']}\n\n"
        
        # Crear placeholder de imagen (archivo de texto)
        placeholder_path = os.path.join(screenshot_dir, f"{page['filename']}.info")
        with open(placeholder_path, 'w', encoding='utf-8') as f:
            f.write(f"PLACEHOLDER: {page['title']}\n")
            f.write(f"Descripción: {page['description']}\n")
            f.write(f"Archivo esperado: {page['filename']}\n")
            f.write(f"\nPara generar este screenshot automáticamente:\n")
            f.write(f"1. Asegurar que la aplicación esté corriendo en localhost:3000\n")
            f.write(f"2. Instalar ChromeDriver\n")
            f.write(f"3. Ejecutar: python capture_screenshots.py\n")
    
    info_content += """
## Cómo generar screenshots automáticos:

1. **Instalar ChromeDriver:**
   - Descargar desde: https://chromedriver.chromium.org/
   - Agregar al PATH del sistema

2. **Ejecutar aplicación:**
   ```bash
   cd inmuebles-web
   npm run dev
   ```

3. **Capturar screenshots:**
   ```bash
   python capture_screenshots.py
   ```

## Alternativa manual:

Si el script automático no funciona, puedes tomar screenshots manualmente:
1. Navegar a cada URL en el navegador
2. Guardar screenshot con el nombre indicado
3. Colocar en este directorio

## URLs principales:

- http://localhost:3000/login
- http://localhost:3000/financial-agent
- http://localhost:3000/financial-agent/movements
- http://localhost:3000/financial-agent/analytics
- http://localhost:3000/financial-agent/contracts
- http://localhost:3000/financial-agent/rules
- http://localhost:3000/financial-agent/integrations
- http://localhost:3000/financial-agent/integrations/bankinter
- http://localhost:3000/financial-agent/mortgage-calculator
- http://localhost:3000/financial-agent/euribor
- http://localhost:3000/financial-agent/documents
- http://localhost:3000/financial-agent/notifications
- http://localhost:3000/financial-agent/tax-assistant
- http://localhost:3000/financial-agent/smart-classifier
- http://localhost:3000/financial-agent/property/1
- http://localhost:3000/financial-agent/property/1/reports
- http://localhost:3000/financial-agent/property/1/mortgage
- http://localhost:3000/financial-agent/property/1/rules
"""
    
    info_file = os.path.join(screenshot_dir, "README_SCREENSHOTS.md")
    with open(info_file, 'w', encoding='utf-8') as f:
        f.write(info_content)
    
    print(f"Directorio de screenshots creado: {screenshot_dir}")
    print(f"Información guardada en: README_SCREENSHOTS.md")
    print(f"Placeholders creados para {len(pages_info)} páginas")
    
    return True

def update_documentation_with_screenshot_info():
    """Actualizar documentación con información sobre screenshots"""
    
    try:
        from docx import Document
        from docx.shared import Inches
        
        # Cargar documento principal
        doc_path = r"C:\Users\davsa\inmuebles\backend\Documentacion_Aplicacion_Inmuebles.docx"
        doc = Document(doc_path)
        
        # Agregar nueva sección al final
        doc.add_page_break()
        doc.add_heading('ANEXO: CAPTURAS DE PANTALLA DETALLADAS', level=1)
        
        doc.add_paragraph(
            'A continuación se detalla la información sobre las capturas de pantalla '
            'de cada funcionalidad principal del sistema:'
        )
        
        pages_info = [
            ('Login', 'Sistema de autenticación con validación de credenciales'),
            ('Dashboard Principal', 'Hub central con resumen de propiedades y métricas'),
            ('Gestión de Movimientos', 'CRUD completo con filtros y clasificación inteligente'),
            ('Analytics Dashboard', 'KPIs avanzados y visualizaciones de rentabilidad'),
            ('Gestión de Contratos', 'Administración de inquilinos y alquileres'),
            ('Reglas de Clasificación', 'Sistema de automatización para categorización'),
            ('Centro de Integraciones', 'Gestión de conexiones bancarias y servicios'),
            ('Integración Bankinter', 'Configuración PSD2 y sincronización automática'),
            ('Calculadora Hipotecaria', 'Simulaciones y cálculo de cuotas variables'),
            ('Gestión Euribor', 'Seguimiento histórico de tipos de interés'),
            ('Gestor de Documentos', 'Sistema de archivos por propiedad'),
            ('Centro de Notificaciones', 'Alertas y recordatorios del sistema'),
            ('Asistente Fiscal', 'Herramientas para optimización fiscal'),
            ('Clasificador Inteligente', 'Sistema de IA para clasificación automática'),
            ('Vista de Propiedad', 'Dashboard específico con métricas detalladas'),
            ('Informes de Propiedad', 'Reportes financieros con análisis mensual'),
            ('Gestión Hipotecaria', 'Administración de hipoteca por propiedad'),
            ('Reglas de Propiedad', 'Configuración específica por propiedad')
        ]
        
        for i, (title, description) in enumerate(pages_info, 1):
            doc.add_heading(f'{i}. {title}', level=2)
            doc.add_paragraph(description)
            doc.add_paragraph(f'📷 [SCREENSHOT: {title.lower().replace(" ", "_")}.png]')
            doc.add_paragraph('')
        
        # Agregar instrucciones para generar screenshots
        doc.add_heading('Instrucciones para Generar Screenshots', level=2)
        
        instructions = [
            '1. Asegurar que la aplicación esté funcionando en localhost:3000',
            '2. Instalar ChromeDriver y agregarlo al PATH del sistema',
            '3. Ejecutar el comando: python capture_screenshots.py',
            '4. Los screenshots se guardarán automáticamente en la carpeta screenshots/',
            '5. Alternativamente, tomar capturas manuales navegando a cada URL'
        ]
        
        for instruction in instructions:
            doc.add_paragraph(instruction)
        
        # Guardar documento actualizado
        doc.save(doc_path)
        print(f"Documentación actualizada con información de screenshots: {doc_path}")
        
        return True
        
    except Exception as e:
        print(f"Error actualizando documentación: {str(e)}")
        return False

if __name__ == "__main__":
    print("Creando estructura para screenshots...")
    
    if create_screenshot_placeholders():
        print("Estructura de screenshots creada exitosamente")
    
    if update_documentation_with_screenshot_info():
        print("Documentación actualizada con información de screenshots")
    
    print("\nNOTA: Para generar screenshots automáticos:")
    print("1. Instalar ChromeDriver: https://chromedriver.chromium.org/")
    print("2. Ejecutar: python capture_screenshots.py")
    print("3. O tomar capturas manuales usando las URLs listadas")
    
    print(f"\nArchivos creados:")
    print(f"- screenshots/README_SCREENSHOTS.md - Guía completa")
    print(f"- screenshots/*.info - Placeholders informativos")
    print(f"- Documentacion_Aplicacion_Inmuebles.docx - Actualizada")